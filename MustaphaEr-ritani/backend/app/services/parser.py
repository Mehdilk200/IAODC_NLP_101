"""
CV / Document Parser Service.
Extracts raw text from PDF and DOCX files.
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using pdfplumber (primary)
    with PyMuPDF as fallback.
    """
    text = ""

    # Primary: pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text.strip()
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, trying PyMuPDF...")

    # Fallback: PyMuPDF (fitz)
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return text.strip()
    except Exception as e:
        logger.error(f"PyMuPDF also failed: {e}")
        raise ValueError("Could not extract text from PDF file.")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError("Could not extract text from DOCX file.")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Route extraction based on file extension.
    Supports PDF and DOCX.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Use PDF or DOCX.")
